from pathlib import Path
import re
from docx import Document
from IMAGES.image_parser import parse_image
import io
from PIL import Image

try:
    import pytesseract
except Exception:
    pytesseract = None

DEBUG = True


# =========================================================
# HELPERS
# =========================================================
def _clean_text(text):
    if text is None:
        return ""

    text = str(text).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()

def _extract_front_matter(text):
    text = _clean_text(text)

    if not text.startswith("---\n"):
        return None, text

    parts = text.split("\n---", 1)
    if len(parts) != 2:
        return None, text

    front_matter = parts[0].strip()
    remaining = parts[1].strip()

    return front_matter, remaining


def _split_markdown_lines(text):
    return text.replace("\r\n", "\n").replace("\r", "\n").split("\n")


def _is_heading_line(line):
    line = line.strip()
    return bool(re.match(r"^#{1,6}\s+", line))


def _is_bullet_line(line):
    line = line.strip()
    return bool(re.match(r"^[-*•]\s+", line))


def _is_numbered_line(line):
    line = line.strip()
    return bool(re.match(r"^\d+[\.\)]\s+", line))

def _detect_block_type(block_text):
    text = block_text.strip()

    if not text:
        return "empty"

    # obsidian / markdown image placeholder
    if re.match(r"^!\[\[.*\]\]$", text):
        return "image_placeholder"

    # fenced code block
    if text.startswith("```") and text.endswith("```"):
        return "code"

    # numbered step / procedural line
    if re.match(r"^\d+[\.\)]\s+", text):
        return "bullet"

    # bullet line
    if re.match(r"^[-*•]\s+", text):
        return "bullet"

    # path / command-ish single line
    if (
        "\\" in text
        or "/" in text
        or text.lower().endswith((".bat", ".sh", ".ps1", ".cmd"))
        or re.match(r"^[A-Za-z]:\\", text)
    ):
        if "\n" not in text and len(text) <= 200:
            return "code"

    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]

    if len(lines) == 1:
        line = lines[0]

        if len(line) <= 80 and not line.endswith((".", "!", "?", ":", ";")):
            words = line.split()

            if words:
                alpha_words = [w for w in words if any(c.isalpha() for c in w)]
                if alpha_words:
                    titlecase_like = sum(1 for w in alpha_words if w[:1].isupper())
                    ratio = titlecase_like / len(alpha_words)
                    if ratio >= 0.6:
                        return "heading"

            if len(words) <= 6:
                return "heading"

    return "paragraph"

def _split_inline_image_placeholders(line):
    parts = re.split(r'(!\[\[[^\]]+\]\])', line)
    return [p.strip() for p in parts if p and p.strip()]

def _extract_placeholder_target(text):
    """
    From ![[image.png]] or ![[image.png|400]] return image.png
    """
    m = re.match(r"^!\[\[(.*)\]\]$", (text or "").strip())
    if not m:
        return None

    inner = m.group(1).strip()
    if not inner:
        return None

    # remove Obsidian display/options part after |
    inner = inner.split("|", 1)[0].strip()

    return inner or None

def _resolve_embedded_image_path(doc_path, target, template_config=None):
    """
    Resolve embedded image reference safely.
    Supports:
    - absolute path
    - relative to current doc folder
    - search under doc folder tree
    - optional asset_search_roots from template_config
    """
    template_config = template_config or {}

    if not target:
        return None

    doc_path = Path(doc_path)
    target_path = Path(target)

    # absolute path
    if target_path.is_absolute() and target_path.exists():
        return target_path

    # relative to current doc folder
    candidate = (doc_path.parent / target_path).resolve()
    if candidate.exists():
        return candidate

    # search under current doc folder tree
    matches = list(doc_path.parent.rglob(target_path.name))
    if matches:
        return matches[0]

    # search extra asset roots
    asset_roots = template_config.get("asset_search_roots") or []
    for root in asset_roots:
        try:
            root_path = Path(root)
            if not root_path.exists():
                continue

            # relative path under asset root
            candidate = (root_path / target_path).resolve()
            if candidate.exists():
                return candidate

            # filename search under asset root
            matches = list(root_path.rglob(target_path.name))
            if matches:
                return matches[0]
        except Exception:
            continue

    return None

def _enrich_blocks_with_embedded_images(blocks, doc_path, template_config=None):
    """
    Preserve all existing blocks.
    If an image placeholder resolves successfully, append an OCR-derived block after it.
    """
    template_config = template_config or {}
    enriched = []

    next_block_id = max((b.get("block_id", 0) for b in blocks), default=0) + 1

    for block in blocks:
        enriched.append(block)

        if block.get("block_type") != "image_placeholder":
            continue

        target = _extract_placeholder_target(block.get("text"))
        resolved_path = _resolve_embedded_image_path(doc_path, target, template_config=template_config)

        if not resolved_path:
            continue

        try:
            parsed_image = parse_image(
                file_path=resolved_path,
                template_config={"enable_ocr": True},
            )

            ocr_text = (
                (parsed_image.get("content") or {}).get("ocr_text") or ""
            ).strip()

            if not ocr_text:
                ocr_text = "[Embedded DOCX image present but no OCR text extracted]"

            enriched.append({
                "block_id": next_block_id,
                "block_type": "embedded_image_text",
                "text": ocr_text,
                "image_target": target,
                "image_path": str(resolved_path),
                "image_doc_type": parsed_image.get("doc_type"),
                "image_mode": parsed_image.get("image_mode"),
            })
            next_block_id += 1

        except Exception:
            # keep existing behavior unchanged on failure
            continue

    return enriched


def _split_text_into_blocks(text):
    text = _clean_text(text)
    if not text:
        return []

    blocks = []
    block_id = 1

    front_matter, remaining = _extract_front_matter(text)

    if front_matter:
        blocks.append({
            "block_id": block_id,
            "block_type": "front_matter",
            "text": front_matter
        })
        block_id += 1

    lines = _split_markdown_lines(remaining)

    current_paragraph = []

    def flush_paragraph():
        nonlocal block_id, current_paragraph, blocks
        if current_paragraph:
            para_text = _clean_text("\n".join(current_paragraph))
            if para_text:
                blocks.append({
                    "block_id": block_id,
                    "block_type": _detect_block_type(para_text),
                    "text": para_text
                })
                block_id += 1
            current_paragraph = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if _is_heading_line(stripped):
            flush_paragraph()
            blocks.append({
                "block_id": block_id,
                "block_type": "heading",
                "text": stripped
            })
            block_id += 1
            continue

        if _is_bullet_line(stripped) or _is_numbered_line(stripped):
            flush_paragraph()

            parts = _split_inline_image_placeholders(stripped)

            for part in parts:
                part = part.strip()
                if not part:
                    continue

                part_type = _detect_block_type(part)

                if part_type == "image_placeholder":
                    blocks.append({
                        "block_id": block_id,
                        "block_type": "image_placeholder",
                        "text": part
                    })
                else:
                    blocks.append({
                        "block_id": block_id,
                        "block_type": "bullet",
                        "text": part
                    })

                block_id += 1

            continue

        parts = _split_inline_image_placeholders(stripped)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            if _detect_block_type(part) == "image_placeholder":
                flush_paragraph()
                blocks.append({
                    "block_id": block_id,
                    "block_type": "image_placeholder",
                    "text": part
                })
                block_id += 1
            else:
                current_paragraph.append(part)

    flush_paragraph()

    return blocks

def _ocr_pil_image(img):
    if pytesseract is None:
        return ""

    try:
        gray = img.convert("L")
        text = pytesseract.image_to_string(gray, config="--psm 6")
        return _clean_text(text)
    except Exception:
        return ""


def _extract_docx_embedded_image_blocks(path, start_block_id=1):
    """
    Phase 1:
    extract embedded images from DOCX relationships and OCR them.
    We append OCR blocks after text parsing, without trying to place them exactly inline yet.
    """
    blocks = []
    block_id = start_block_id

    try:
        doc = Document(path)

        for rel in doc.part.rels.values():
            reltype = str(getattr(rel, "reltype", "") or "")
            if "image" not in reltype:
                continue

            try:
                image_blob = rel.target_part.blob
                img = Image.open(io.BytesIO(image_blob))
                ocr_text = _ocr_pil_image(img)

                if not ocr_text:
                    continue

                image_name = getattr(rel.target_part, "partname", None)
                image_name = Path(str(image_name)).name if image_name else f"docx_image_{block_id}.png"

                if DEBUG:
                    print(f"[DOCX IMAGE OCR] {image_name} -> {ocr_text[:120]}")

                blocks.append({
                    "block_id": block_id,
                    "block_type": "embedded_image_text",
                    "text": ocr_text,
                    "image_target": image_name,
                    "image_path": str(path),
                    "image_doc_type": "reference",
                    "image_mode": "embedded_docx_image",
                })
                block_id += 1

            except Exception:
                continue

    except Exception:
        return []

    return blocks  

def _extract_markdown_related_titles(text):
    matches = re.findall(r"\[\[([^|\]]+)(?:\|[^\]]+)?\]\]", text or "")
    seen = set()
    related = []

    for m in matches:
        title = str(m).strip()
        if title and title not in seen:
            seen.add(title)
            related.append(title)

    return related  


# =========================================================
# TXT / MD / RTF PLACEHOLDER EXTRACTION
# =========================================================
def _read_text_file(path):
    ext = path.suffix.lower()

    if ext in [".txt", ".md", ".rtf"]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        doc = Document(path)

        parts = []

        for para in doc.paragraphs:
            text = _clean_text(para.text)
            if text:
                parts.append(text)

        return "\n\n".join(parts)

    raise ValueError(f"Unsupported doc format: {ext}")

# =========================================================
# MAIN PARSER
# =========================================================
def parse_doc(file_path, template_config=None):
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in [".txt", ".md", ".rtf", ".docx"]:
        raise ValueError(f"Unsupported doc format: {ext}")

    raw_text = _read_text_file(path)
    blocks = _split_text_into_blocks(raw_text)

    if ext == ".docx":
        next_block_id = max((b.get("block_id", 0) for b in blocks), default=0) + 1
        docx_image_blocks = _extract_docx_embedded_image_blocks(path, start_block_id=next_block_id)
        blocks.extend(docx_image_blocks)
    else:
        blocks = _enrich_blocks_with_embedded_images(
            blocks=blocks,
            doc_path=path,
            template_config=template_config or {}
        )

    related_titles = _extract_markdown_related_titles(raw_text) if ext == ".md" else []

    result = {
        "blocks": blocks,
        "text": _clean_text(raw_text),
        "source_file": path.name,
        "source_path": str(path),
        "filetype": "docx" if ext == ".docx" else "doc",
        "related_titles": related_titles,
        "block_count": len(blocks)
    }

    if DEBUG:
        print(f"[DOC PARSER] Loaded {result['block_count']} blocks from {result['source_file']}")
        print(f"[DOCX IMAGE OCR] Extracted {len(blocks)} OCR image blocks from {path.name}")
        if blocks:
            print(f"[DOC PARSER] First block type: {blocks[0]['block_type']}")
            print(f"[DOC PARSER] First block text: {blocks[0]['text'][:200]}")

    return result